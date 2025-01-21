from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import textwrap

class BilingualWordGenerator:
    def __init__(self, output_path):
        self.output_path = Path(output_path).absolute()
        
        # 创建新文档
        self.doc = Document()

    
        
        # 设置页面边距
        sections = self.doc.sections
        for section in sections:
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
        
        # 设置文本参数
        self.english_chars_per_line = 58
        self.chinese_chars_per_line = 18
        
        # 创建两列表格作为主要内容容器
        self.table = self.doc.add_table(rows=1, cols=2)
        self.table.autofit = False
        self.table.allow_autofit = False
        
        # 设置列宽比例（60:40）
        self.set_column_widths()
    
    def set_column_widths(self):
        # 设置表格总宽度
        self.table.width = Inches(7)  # A4纸的可用宽度约为7英寸
        
        # 设置列宽比例
        for row in self.table.rows:
            row.cells[0].width = Inches(4.2)  # 60% 的宽度
            row.cells[1].width = Inches(2.8)  # 40% 的宽度
    

    def add_text_to_cell(self, cell, text, is_english=True):
        """向单元格添加文本，设置字体和格式"""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        
        # 设置字体和大小
        font = run.font
        if is_english:
            font.name = 'Helvetica'
            # 确保英文字体设置
            font.ascii_font = font.hAnsi_font = 'Helvetica'
        else:
            # 明确设置中文字体
            font.name = 'Source Han Serif'#Microsoft YaHei'
            # 专门设置东亚文字字体
            font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Source Han Serif')
        font.size = Pt(12)
    
    def process_text_files(self, english_file, chinese_file):
        try:
            # 读取文本文件
            with open(english_file, 'r', encoding='utf-8') as f:
                english_text = f.read().split('\n')
            
            with open(chinese_file, 'r', encoding='utf-8') as f:
                chinese_text = f.read().split('\n')
            
            # 处理每对文本
            for en, zh in zip(english_text, chinese_text):
                if not en.strip() and not zh.strip():  # 跳过空行
                    continue
                    
                # 添加新行
                row_cells = self.table.add_row().cells
                
                # 对长文本进行换行处理
                en_wrapped = textwrap.fill(en, width=self.english_chars_per_line)
                zh_wrapped = textwrap.fill(zh, width=self.chinese_chars_per_line)
                
                # 添加英文和中文文本到对应的单元格
                self.add_text_to_cell(row_cells[0], en_wrapped, is_english=True)
                self.add_text_to_cell(row_cells[1], zh_wrapped, is_english=False)
            
            # 删除第一个空行
            if len(self.table.rows) > 1:
                tr = self.table._tbl.tr_lst[0]
                tr.getparent().remove(tr)
            
            # 保存文档
            self.doc.save(str(self.output_path))
            print(f"Word file generated successfully: {self.output_path}")
            
        except Exception as e:
            print(f"Error processing files: {e}")
            raise


def main():
    current_dir = Path.cwd()
    output_word = current_dir / 'mars.docx'
    eng_file_path = 'mars.txt'
    ch_file_path = 'mars_ch.txt'
    
    try:
        generator = BilingualWordGenerator(output_word)
        generator.process_text_files(str(current_dir / eng_file_path), str(current_dir / ch_file_path))
        
    except Exception as e:
        print(f"Program execution error: {e}")

if __name__ == '__main__':
    main()